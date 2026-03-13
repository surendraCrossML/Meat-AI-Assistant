import io
import uuid
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from typing import Optional
from sqlalchemy import or_
from app.db.base import get_db
from app.models.document import Document
from app.services.s3_service import (
    upload_file_to_s3,
    generate_presigned_url,
    delete_file_from_s3,
    list_files_in_s3,
    read_file_from_s3,
)
from app.services.gemini_service import (
    extract_keywords_from_query,
    generate_response_from_documents,
)
from app.schemas.schemas import (
    DocumentRead,
    PresignedURLResponse,
    DocumentSearchRequest,
    DocumentQueryRequest,
    DocumentQueryResponse,
    KeywordExtractionResponse,
)

# PDF and DOCX parsing libraries
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


router = APIRouter(prefix="/documents", tags=["Documents"])

# S3 folder — all documents go here
S3_FOLDER = "beef-documents"


# --------------------------------
# Document text extraction helpers
# --------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Parse a PDF file from raw bytes and return its full text content.
    Uses pypdf to extract text from every page.
    """
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")

    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages_text.append(f"[Page {page_num}]\n{text}")

    full_text = "\n\n".join(pages_text)
    print(f"[PDF PARSER] Extracted {len(reader.pages)} page(s), {len(full_text)} characters total.")
    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Parse a DOCX file from raw bytes and return its full text content.
    Concatenates all paragraph text in document order.
    """
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    full_text = "\n".join(paragraphs)
    print(f"[DOCX PARSER] Extracted {len(paragraphs)} paragraph(s), {len(full_text)} characters total.")
    return full_text


def extract_text_from_file(s3_key: str, file_bytes: bytes) -> str:
    """
    Route file bytes to the correct parser based on the S3 key extension.
    - .pdf  → extract_text_from_pdf
    - .docx → extract_text_from_docx
    - other → decode as UTF-8 text
    """
    key_lower = s3_key.lower()
    if key_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif key_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        # Plain text / CSV / XML etc.
        text = file_bytes.decode("utf-8", errors="replace")
        print(f"[TEXT PARSER] Decoded {len(text)} characters as plain text.")
        return text


@router.post(
    "/upload",
    response_model=DocumentRead,
    status_code=status.HTTP_201_CREATED,
    summary="Upload a document to S3 and save metadata in the database",
)
async def upload_document(
    file: UploadFile = File(...),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    # Build a unique S3 key inside the fixed beef-documents folder
    s3_key = f"{S3_FOLDER}/{uuid.uuid4()}-{file.filename}"

    try:
        upload_file_to_s3(file.file, s3_key, file.content_type or "application/octet-stream")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"S3 upload failed: {exc}",
        )

    doc = Document(
        document_name=file.filename,
        document_type=file.content_type or "application/octet-stream",
        document_size=file.size or 0,
        description=description,
        s3_key=s3_key,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


@router.get(
    "/",
    response_model=list[DocumentRead],
    summary="List all documents",
)
def list_documents(
    db: Session = Depends(get_db),
):
    return db.query(Document).all()


@router.get(
    "/{document_id}/download",
    response_model=PresignedURLResponse,
    summary="Get a presigned URL to download a document from S3",
)
def download_document(
    document_id: int,
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == document_id).first()

    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    try:
        url = generate_presigned_url(doc.s3_key)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Could not generate presigned URL: {exc}",
        )

    return {"presigned_url": url, "expires_in_seconds": 3600}


@router.delete(
    "/{document_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete a document from S3 and remove its metadata from the database",
)
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == document_id).first()

    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    try:
        delete_file_from_s3(doc.s3_key)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"S3 delete failed: {exc}",
        )

    db.delete(doc)
    db.commit()


@router.post(
    "/query",
    response_model=DocumentQueryResponse,
    summary="Query documents using natural language with AI keyword extraction",
)
def query_documents_with_ai(
    request: DocumentQueryRequest,
    db: Session = Depends(get_db),
):
    """
    Query documents using natural language.

    This endpoint:
    1. Takes a user query in natural language
    2. Uses Gemini AI to extract relevant keywords (recipes, beef-related terms)
    3. Searches the database for matching documents
    4. Returns matching documents with extraction details

    Request body:
    - **query**: Your question or search intent (e.g., "What are the best beef recipes for a healthy diet?")

    Response includes:
    - Extracted keywords and summaries
    - List of matching documents
    - Summary of the search
    """
    try:
        # Extract keywords using Gemini AI
        keywords_data = extract_keywords_from_query(request.query)

        # Prepare keyword extraction response
        extracted_keywords = KeywordExtractionResponse(**keywords_data)

        # Combine all keywords for comprehensive search
        all_keywords = extracted_keywords.keywords

        # Search documents using extracted keywords
        matching_documents = []

        filters = []
        for keyword in all_keywords:
            term = f"%{keyword}%"
            filters.append(Document.document_name.ilike(term))
            filters.append(Document.description.ilike(term))

        docs = db.query(Document).filter(or_(*filters)).all()   

        # Add unique documents to the matching list
        for doc in docs:
            if not any(d.id == doc.id for d in matching_documents):
                matching_documents.append(doc)

        # Create summary
        summary = f"Found {len(matching_documents)} document(s) matching your search for: {extracted_keywords.query_summary}"

        # Fetch actual file content from S3 for each matched document
        # print(f"\n{'='*60}")
        # print(f"[QUERY] User asked: {request.query}")
        # print(f"[QUERY] Matched {len(matching_documents)} document(s) in DB: {[d.document_name for d in matching_documents]}")
        # print(f"{'='*60}")

        docs_for_llm = []
        for doc in matching_documents:
            print(f"\n[S3 DOWNLOAD] Fetching: '{doc.document_name}' (s3_key={doc.s3_key})")
            try:
                from app.services.s3_service import download_bytes_from_s3
                raw_bytes = download_bytes_from_s3(doc.s3_key)
                print(f"[S3 DOWNLOAD] Downloaded {len(raw_bytes)} bytes for '{doc.document_name}'")
                file_content = extract_text_from_file(doc.s3_key, raw_bytes)
            except Exception as e:
                # Fall back to description if S3 read or parsing fails
                print(f"[S3 DOWNLOAD] ERROR for '{doc.document_name}': {e} — falling back to DB description.")
                file_content = doc.description or ""

            # DEBUG: Show a snippet of what the LLM will actually read
            snippet = file_content[:300].replace("\n", " ")
            # print(f"[LLM INPUT] First 300 chars of '{doc.document_name}': {snippet!r}")
            # print(f"[LLM INPUT] Total content length sent to LLM: {len(file_content)} characters")

            docs_for_llm.append({
                "document_name": doc.document_name,
                "content": file_content,
            })

        # print(f"\n[LLM CALL] Sending {len(docs_for_llm)} document(s) to Gemini for response generation...")
        # Ask Gemini to produce a response based on the documents
        try:
            agent_response = generate_response_from_documents(request.query, docs_for_llm)
        except Exception:
            agent_response = "Unable to generate response from documents."

        return DocumentQueryResponse(
            user_query=request.query,
            extracted_keywords=extracted_keywords,
            matching_documents=matching_documents,
            summary=summary,
            agent_response=agent_response,
        )

    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"AI service error: {str(e)}",
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing query: {str(e)}",
        )