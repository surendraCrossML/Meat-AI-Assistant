"""
text_extractor.py
~~~~~~~~~~~~~~~~~~
Extracts plain text from DOCX and PDF files downloaded from S3.

Strategy
--------
1. Determine the file type from the s3_key extension (or fallback to content
   sniffing).
2. For PDF  → use pdfplumber (accurate, handles most modern PDFs).
3. For DOCX → use python-docx (reads paragraph and table text).
4. For plain text / markdown → decode bytes as UTF-8 directly.
"""

import io
import logging
from pathlib import PurePosixPath

logger = logging.getLogger(__name__)


def _ext(s3_key: str) -> str:
    """Return lowercase file extension from an S3 key, e.g. '.pdf'."""
    return PurePosixPath(s3_key).suffix.lower()


def extract_text_from_bytes(raw_bytes: bytes, s3_key: str) -> str:
    """
    Convert raw binary content from S3 into plain text.

    Args:
        raw_bytes: The raw bytes downloaded from S3.
        s3_key:    The S3 object key (used to determine file type by extension).

    Returns:
        Extracted plain text suitable for sending to an LLM.
        Returns an empty string if extraction fails entirely.
    """
    ext = _ext(s3_key)

    # PDF
    if ext == ".pdf":
        return _extract_pdf(raw_bytes, s3_key)

    # DOCX
    if ext in (".docx", ".doc"):
        return _extract_docx(raw_bytes, s3_key)

    # Plain text / Markdown / CSV
    if ext in (".txt", ".md", ".csv", ".json"):
        return raw_bytes.decode("utf-8", errors="replace")

    # Unknown: try UTF-8 first, fall back to partial decode
    logger.warning(
        "[TextExtractor] Unknown extension '%s' for key %s — trying UTF-8 decode.",
        ext,
        s3_key,
    )
    return raw_bytes.decode("utf-8", errors="replace")


# PDF helper
def _extract_pdf(raw_bytes: bytes, s3_key: str) -> str:
    try:
        import pdfplumber  # imported here so it's optional at module level

        pages_text = []
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text.strip())

        text = "\n\n".join(pages_text)
        logger.info(
            "[TextExtractor] PDF extracted %d chars from %s (%d pages)",
            len(text),
            s3_key,
            len(pages_text),
        )
        return text

    except Exception as exc:
        logger.error("[TextExtractor] PDF extraction failed for %s: %s", s3_key, exc)
        return ""


# DOCX helper
def _extract_docx(raw_bytes: bytes, s3_key: str) -> str:
    try:
        import docx as python_docx  # python-docx package

        doc = python_docx.Document(io.BytesIO(raw_bytes))

        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables (flatten cell text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        logger.info(
            "[TextExtractor] DOCX extracted %d chars from %s",
            len(text),
            s3_key,
        )
        return text

    except Exception as exc:
        logger.error("[TextExtractor] DOCX extraction failed for %s: %s", s3_key, exc)
        return ""
